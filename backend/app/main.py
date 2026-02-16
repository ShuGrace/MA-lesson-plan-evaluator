from fastapi import FastAPI, Depends, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from app.db.database import Database, init_database
from typing import List, Optional
from pydantic import BaseModel
import traceback
from contextlib import asynccontextmanager
from io import BytesIO
import json
import time
import asyncio
from functools import wraps
import logging

# ============================================================
# Logging Setup
# ============================================================
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger("lesson-evaluator")


# ============================================================
# Performance Monitoring Decorator
# ============================================================
def timing_decorator(func):
    @wraps(func)
    async def wrapper(*args, **kwargs):
        start_time = time.time()
        try:
            result = await func(*args, **kwargs)
            elapsed = time.time() - start_time
            logger.info(f"{func.__name__} executed in {elapsed:.2f}s")
            return result
        except Exception as e:
            elapsed = time.time() - start_time
            logger.error(f"{func.__name__} failed after {elapsed:.2f}s: {e}")
            raise
    return wrapper


# ============================================================
# Score Colour Helper
# ============================================================
def score_color(score: Optional[int]) -> str:
    if score is None:
        return "grey"
    try:
        return "red" if score < 60 else "orange" if score < 80 else "green"
    except Exception:
        return "grey"


# ============================================================
# Optional File Processing Libraries
# ============================================================
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed")

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not installed")


# ============================================================
# App-level Imports
# ============================================================
from app.config import (
    API_MODE,
    ENABLE_DEEPSEEK,
    ENABLE_CLAUDE,
    ENABLE_GPT,
    API_TIMEOUT,
    API_MAX_RETRIES,
    API_RETRY_DELAY,
)
from app.services.llm_client import LLMClient
from app.services.framework_loader import get_framework_loader
from app.utils.evaluation_helpers import (
    extract_score_from_response,
    extract_recommendations_from_response,
    extract_strengths_from_response,
    extract_areas_for_improvement_from_response,
    parse_json_response,
    calculate_weighted_score,
    merge_and_deduplicate_recommendations,
)

# Initialise framework loader (singleton)
framework_loader = get_framework_loader()


# ============================================================
# Application Lifespan
# ============================================================
@asynccontextmanager
async def lifespan(app: FastAPI):
    try:
        logger.info("Starting application...")
        init_database(reset=False)
        logger.info("Database initialised successfully")
        logger.info("Framework will be loaded on first use")
    except Exception as e:
        logger.warning(f"Initialisation warning: {e}")
        traceback.print_exc()

    yield  # Application is now accepting requests

    logger.info("Application shutdown")


app = FastAPI(title="Lesson Plan Evaluator API v3.0", lifespan=lifespan)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],
    max_age=3600,
)


# ============================================================
# Pydantic Models
# ============================================================
class EvaluationCreate(BaseModel):
    lesson_plan_text: str
    lesson_plan_title: Optional[str] = None
    grade_level: Optional[str] = None
    subject_area: Optional[str] = None
    provider: Optional[str] = "gpt"


class EvaluationUpdate(BaseModel):
    place_based_score: Optional[int] = None
    cultural_score: Optional[int] = None
    overall_score: Optional[int] = None
    status: Optional[str] = None
    agent_responses: Optional[List[dict]] = None
    debate_transcript: Optional[dict] = None
    recommendations: Optional[List[str]] = None


class ImproveLessonRequest(BaseModel):
    original_lesson: str
    lesson_title: str
    grade_level: Optional[str] = None
    subject_area: Optional[str] = None
    recommendations: List[str]
    scores: dict
    remove_numbering: Optional[bool] = False


class ConvertToWordRequest(BaseModel):
    content: str
    filename: Optional[str] = "Improved_Lesson_Plan.docx"
    title: Optional[str] = "Improved Lesson Plan"


# ============================================================
# Database Dependency
# ============================================================
def get_db():
    db = Database()
    db.connect()
    try:
        yield db
    finally:
        db.close()


# ============================================================
# File Extraction Helpers
# ============================================================
def extract_text_from_docx(file_bytes: bytes) -> str:
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=501, detail="DOCX support not available")
    try:
        doc = docx.Document(BytesIO(file_bytes))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read DOCX: {e}")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    if not PDF_AVAILABLE:
        raise HTTPException(status_code=501, detail="PDF support not available")
    try:
        reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read PDF: {e}")


# ============================================================
# Lesson Plan Validation
# ============================================================
def validate_lesson_format(lesson_text: str) -> dict:
    """
    Validate the generated lesson plan format quality.
    Detects structured formats, Python list syntax, etc.
    """
    issues = []
    warnings = []

    # Detect numbered section format
    if "1.1" in lesson_text or "1.2" in lesson_text or "2.1" in lesson_text:
        issues.append("Contains numbered sections (1.1, 1.2, 2.1, etc.)")

    # Detect Python list syntax
    if "['Understanding" in lesson_text or '["Understanding' in lesson_text:
        issues.append("Contains Python list syntax")

    # Detect JSON format
    if lesson_text.strip().startswith("{") and '"knowledge":' in lesson_text:
        issues.append("Appears to be JSON format instead of narrative")

    # Length checks
    if len(lesson_text) < 1000:
        issues.append("Too short (less than 1000 chars)")
    elif len(lesson_text) < 1500:
        warnings.append("Relatively short (less than 1500 chars)")

    # Te Reo content
    te_reo_terms = [
        "Te Reo", "Māori", "Kia ora", "whānau", "mana",
        "kaitiakitanga", "whanaungatanga", "whakapapa",
    ]
    has_te_reo = any(term in lesson_text for term in te_reo_terms)
    if not has_te_reo:
        warnings.append("No Te Reo Māori terms detected")

    # Specific places
    places = [
        "Auckland", "Wellington", "Canterbury", "Ōrākei", "Te Papa",
        "Museum", "Marae", "Waitangi",
    ]
    has_specific_places = any(place in lesson_text for place in places)
    if not has_specific_places:
        warnings.append("No specific local places named")

    # Critical questions
    critical_phrases = [
        "whose stories", "whose voices", "different perspectives",
        "why might different", "how has",
    ]
    has_critical_questions = any(q in lesson_text.lower() for q in critical_phrases)

    return {
        "valid": len(issues) == 0,
        "issues": issues,
        "warnings": warnings,
        "char_count": len(lesson_text),
        "word_count": len(lesson_text.split()),
        "has_te_reo": has_te_reo,
        "has_specific_places": has_specific_places,
        "has_critical_questions": has_critical_questions,
    }


# ============================================================
# Build Analysis Structure Helper
# ============================================================
def build_analysis_structure(
    dimension_key,
    score,
    response_text,
    recommendations=None,
    strengths=None,
    areas_for_improvement=None,
    gaps=None,
    cultural_elements=None,
):
    """
    Build the analysis structure the frontend expects — Framework v3.0.

    Key distinction:
    - strengths: what was done well
    - areas_for_improvement: problem statements (what needs work)
    - recommendations: specific actionable suggestions (solutions)
    """
    return {
        dimension_key: {
            "score": score,
            "response_text": response_text or "",
            "strengths": strengths or [],
            "areas_for_improvement": areas_for_improvement or [],
            "gaps": gaps or [],
            "recommendations": recommendations or [],
            "cultural_elements_present": cultural_elements or [],
            "summary": (response_text[:500] if response_text else ""),
        }
    }


# ============================================================
# Prompt Loading Helper
# ============================================================
def _load_improvement_prompt(
    title: str,
    grade_level: str,
    subject_area: str,
    text: str,
    scores: dict,
    all_strengths: list,
    all_areas: list,
    recommendations: list,
) -> str:
    """
    Build the improvement prompt for Claude.
    Extracted from the endpoint to keep the function body manageable.
    """
    strengths_text = "\n".join(
        f"- {s}" for s in all_strengths[:6]
    ) if all_strengths else "(Limited strengths identified)"

    areas_text = "\n".join(
        f"- {a}" for a in all_areas[:6]
    ) if all_areas else "(See recommendations)"

    recs_text = "\n".join(
        f"{i + 1}. {rec}" for i, rec in enumerate(recommendations[:10])
    )

    place_based_score = scores.get("place_based_learning", 0)
    cultural_score = scores.get("cultural_responsiveness_integrated", 0)
    critical_pedagogy_score = scores.get("critical_pedagogy", 0)
    lesson_design_score = scores.get("lesson_design_quality", 0)
    overall_score = scores.get("overall", 0)

    return f"""You are a highly experienced educator in Aotearoa New Zealand.
Your task is to significantly improve this lesson plan based on detailed evaluation feedback.

CRITICAL OUTPUT FORMAT REQUIREMENT:
You MUST write as a narrative lesson plan document, NOT as code or structured data.
- DO NOT use Python lists like ['item1', 'item2', 'item3']
- DO NOT use numbered sections like 1.1, 1.2, 2.1
- DO NOT format output as JSON or dictionary
- DO write in flowing narrative paragraphs like a real lesson plan
- DO write naturally as if you're a teacher writing for other teachers

ORIGINAL LESSON:
Title: {title}
Grade: {grade_level or 'Not specified'}
Subject: {subject_area or 'Not specified'}

Original Plan (excerpt):
{text[:3000]}

EVALUATION FEEDBACK:
Scores:
- Place-based: {place_based_score}/100
- Cultural Responsiveness: {cultural_score}/100
- Critical Pedagogy: {critical_pedagogy_score}/100
- Lesson Design: {lesson_design_score}/100
- Overall: {overall_score}/100

Strengths:
{strengths_text}

Areas for Improvement:
{areas_text}

Recommendations:
{recs_text}

REQUIREMENTS:
1. Include at least 3 Te Reo Maori terms with correct macrons (a, e, i, o, u)
2. Name 1-2 specific local places (e.g., Orakei Marae, Auckland Museum, Te Papa)
3. Include 1-2 critical thinking questions (e.g., "Whose stories do we usually hear?")
4. Provide 2-4 activities described in narrative paragraph form
5. Include an assessment rubric table
6. List 5-8 specific resources with titles and URLs
7. Target 1200-2000 words

Include a cultural disclaimer at the end:
"This AI-generated lesson plan requires review and adaptation before use. Please consult with local iwi and cultural advisors for Maori content."

Write the improved lesson plan now in flowing narrative paragraphs, starting with:
**IMPROVED LESSON PLAN: {title.upper()}**
"""


# ============================================================
# API Endpoints
# ============================================================

@app.get("/")
async def root():
    """API root — system status and theoretical framework info."""
    framework = framework_loader.load_theoretical_framework()
    agent_design = framework_loader.load_agent_design()

    return {
        "message": "Lesson Plan Evaluator API with Theoretical Framework v3.0",
        "mode": API_MODE,
        "database": "connected",
        "status": "operational",
        "theoretical_framework": {
            "name": framework.get("framework_metadata", {}).get("name", "Default"),
            "version": framework.get("framework_metadata", {}).get("version", "3.0"),
            "dimensions": list(framework.get("dimensions", {}).keys()),
            "agents_configured": len(agent_design.get("agents", {})),
        },
        "features": {
            "file_upload": {"word": DOCX_AVAILABLE, "pdf": PDF_AVAILABLE}
        },
        "api_status": {
            "deepseek": "enabled" if ENABLE_DEEPSEEK else "disabled",
            "claude": "enabled" if ENABLE_CLAUDE else "disabled",
            "gpt": "enabled" if ENABLE_GPT else "disabled",
        },
    }


@app.get("/api/framework")
async def get_framework_info():
    """Return detailed framework information."""
    framework = framework_loader.load_theoretical_framework()
    agent_design = framework_loader.load_agent_design()
    weights = framework_loader.get_scoring_weights()

    dimensions_summary = {}
    for dim_code, dim_info in framework.get("dimensions", {}).items():
        dimensions_summary[dim_code] = {
            "label": dim_info.get("label", ""),
            "definition": dim_info.get("definition", ""),
            "indicator_count": len(dim_info.get("indicators", [])),
            "weight": weights.get(dim_code, 0.0),
        }

    agents_summary = {}
    for _agent_id, agent_info in agent_design.get("agents", {}).items():
        agents_summary[agent_info["name"]] = {
            "role": agent_info.get("role", ""),
            "dimensions": agent_info.get("assigned_dimensions", []),
        }

    return {
        "status": "success",
        "framework_metadata": framework.get("framework_metadata", {}),
        "dimensions": dimensions_summary,
        "agents": agents_summary,
        "composite_scoring": {
            "method": framework.get("composite_scoring", {}).get(
                "method", "weighted_average"
            ),
            "weights": weights,
        },
    }


@app.get("/api/framework/dimension/{dimension_code}")
async def get_dimension_details(dimension_code: str):
    """Return detailed indicator information for a specific dimension."""
    framework = framework_loader.load_theoretical_framework()
    dimensions = framework.get("dimensions", {})

    if dimension_code not in dimensions:
        raise HTTPException(
            status_code=404,
            detail=f"Dimension '{dimension_code}' not found",
        )

    dim = dimensions[dimension_code]
    indicators = [
        {
            "code": ind.get("code", ""),
            "name": ind.get("name", ""),
            "definition": ind.get("definition", ""),
            "source": ind.get("source", ""),
        }
        for ind in dim.get("indicators", [])
    ]

    return {
        "status": "success",
        "code": dimension_code,
        "label": dim.get("label", ""),
        "definition": dim.get("definition", ""),
        "theoretical_foundation": dim.get("theoretical_foundation", {}),
        "indicators": indicators,
        "scoring_rubric": dim.get("scoring_rubric", {}),
    }


@app.post("/api/upload-file")
async def upload_lesson_plan_file(file: UploadFile = File(...)):
    """Upload a DOCX or PDF file and extract text."""
    try:
        file_bytes = await file.read()

        if file.filename.endswith(".docx"):
            text = extract_text_from_docx(file_bytes)
        elif file.filename.endswith(".pdf"):
            text = extract_text_from_pdf(file_bytes)
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Use .docx or .pdf",
            )

        return {
            "status": "success",
            "filename": file.filename,
            "text": text,
            "length": len(text),
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"File processing failed: {e}")


@app.post("/api/extract-text")
async def extract_text_from_file(file: UploadFile = File(...)):
    """Extract text content from an uploaded PDF or DOCX file."""
    try:
        logger.info(f"Processing uploaded file: {file.filename} ({file.content_type})")

        file_bytes = await file.read()
        logger.info(f"File size: {len(file_bytes) / 1024:.1f} KB")

        text = ""
        metadata = {}

        is_docx = file.filename.endswith(".docx") or file.content_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        is_pdf = file.filename.endswith(".pdf") or file.content_type == "application/pdf"

        if is_docx:
            if not DOCX_AVAILABLE:
                raise HTTPException(
                    status_code=500,
                    detail="DOCX processing not available. Install python-docx.",
                )

            text = extract_text_from_docx(file_bytes)

            # Try to extract metadata
            try:
                doc = docx.Document(BytesIO(file_bytes))
                if doc.paragraphs and doc.paragraphs[0].text.strip():
                    first_para = doc.paragraphs[0].text.strip()
                    if len(first_para) < 100 and not first_para.endswith("."):
                        metadata["title"] = first_para
                if hasattr(doc.core_properties, "title") and doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title
            except Exception as meta_err:
                logger.warning(f"Could not extract metadata: {meta_err}")

        elif is_pdf:
            if not PDF_AVAILABLE:
                raise HTTPException(
                    status_code=500,
                    detail="PDF processing not available. Install PyPDF2.",
                )
            text = extract_text_from_pdf(file_bytes)

        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported file type. Only PDF and Word (.docx, .doc) files are supported.",
            )

        if not text or len(text.strip()) < 10:
            raise HTTPException(
                status_code=400,
                detail="Could not extract sufficient text. File may be empty or corrupted.",
            )

        logger.info(f"Extracted {len(text)} characters from {file.filename}")

        return {
            "status": "success",
            "filename": file.filename,
            "text": text.strip(),
            "length": len(text),
            "metadata": metadata,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"File extraction error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"File processing failed: {e}")


@app.get("/api/health")
async def health_check():
    """Health check endpoint."""
    return {
        "status": "ok",
        "message": "Service is healthy",
        "framework_loaded": framework_loader._framework is not None,
        "framework_version": "3.0",
        "api_config": {
            "deepseek_enabled": ENABLE_DEEPSEEK,
            "claude_enabled": ENABLE_CLAUDE,
            "gpt_enabled": ENABLE_GPT,
            "timeout": API_TIMEOUT,
            "max_retries": API_MAX_RETRIES,
            "retry_delay": API_RETRY_DELAY,
        },
    }


# ============================================================
# Main Evaluation Endpoint
# ============================================================
@app.post("/api/evaluate")
@app.post("/api/evaluate/lesson")
@timing_decorator
async def evaluate_lesson_plan(
    request: EvaluationCreate,
    db: Database = Depends(get_db),
):
    """
    Evaluate a lesson plan — Framework v3.0.
    4 dimensions: Place-based, Cultural Responsiveness (Integrated),
                  Critical Pedagogy, Lesson Design Quality.
    """
    logger.info("=" * 60)
    logger.info("New Evaluation Request (Framework v3.0)")
    logger.info("=" * 60)

    text = request.lesson_plan_text.strip()
    title = request.lesson_plan_title or "Untitled Lesson Plan"
    grade_level = request.grade_level
    subject_area = request.subject_area
    provider = request.provider or "gpt"

    if not text:
        raise HTTPException(status_code=400, detail="Lesson plan text cannot be empty")
    if provider not in ("gpt", "claude"):
        raise HTTPException(status_code=400, detail="Provider must be 'gpt' or 'claude'")

    logger.info(f"Title: {title}")
    logger.info(f"Grade: {grade_level}, Subject: {subject_area}")
    logger.info(f"Length: {len(text)} chars, Provider: {provider.upper()}")

    llm_client = LLMClient()

    # Load framework v3.0 prompts
    logger.info("Loading framework v3.0 prompts...")
    deepseek_prompt_template = framework_loader.load_prompt("deepseek")
    claude_prompt_template = framework_loader.load_prompt("claude")
    gpt_critical_prompt_template = framework_loader.load_prompt("gpt_critical")
    gpt_design_prompt_template = framework_loader.load_prompt("gpt_design")

    # Initialise result variables
    agent_responses = []
    place_based_score = 0
    cultural_score = 0
    critical_pedagogy_score = 0
    lesson_design_score = 0
    overall_score = 0
    lesson_plan_text = text  # BUG FIX: default to original text, not empty string
    recommendations = []
    eval_id = None

    if API_MODE == "real":
        logger.info(f"REAL API mode | DeepSeek={ENABLE_DEEPSEEK}, Claude={ENABLE_CLAUDE}, GPT={ENABLE_GPT}")
        logger.info(f"Timeout={API_TIMEOUT}s, Retries={API_MAX_RETRIES}, Delay={API_RETRY_DELAY}s")

        try:
            # Determine which LLM to call based on provider
            llm_name = "chatgpt" if provider == "gpt" else "claude"
            model_name = "gpt-4o" if provider == "gpt" else "claude-sonnet-4-20250514"

            # ── AGENT 1: Place-Based Learning ──
            logger.info(f"Agent 1/{provider.upper()}: Evaluating Place-Based Learning...")
            pbl_prompt = deepseek_prompt_template.format(lesson_plan_text=text)
            pbl_response = await llm_client.call(
                llm_name, pbl_prompt, agent_name=f"{provider.upper()}-PlaceBased"
            )

            place_based_score = extract_score_from_response(pbl_response, "place_based")
            pbl_recommendations = extract_recommendations_from_response(pbl_response)
            pbl_strengths = extract_strengths_from_response(pbl_response)
            pbl_areas = extract_areas_for_improvement_from_response(pbl_response)
            logger.info(f"Place-Based Score: {place_based_score}/100")

            # ── AGENT 2: Cultural Responsiveness & Māori Perspectives ──
            logger.info(f"Agent 2/{provider.upper()}: Evaluating Cultural Responsiveness...")
            crmp_prompt = claude_prompt_template.format(lesson_plan_text=text)
            crmp_response = await llm_client.call(
                llm_name, crmp_prompt, agent_name=f"{provider.upper()}-Cultural"
            )

            cultural_score = extract_score_from_response(crmp_response, "cultural")
            crmp_recommendations = extract_recommendations_from_response(crmp_response)
            crmp_strengths = extract_strengths_from_response(crmp_response)
            crmp_areas = extract_areas_for_improvement_from_response(crmp_response)
            logger.info(f"Cultural Responsiveness Score: {cultural_score}/100")

            # ── AGENT 3: Critical Pedagogy ──
            logger.info(f"Agent 3/{provider.upper()}: Evaluating Critical Pedagogy...")
            cp_prompt = gpt_critical_prompt_template.format(lesson_plan_text=text)
            cp_response = await llm_client.call(
                llm_name, cp_prompt, agent_name=f"{provider.upper()}-Critical"
            )

            critical_pedagogy_score = extract_score_from_response(cp_response, "critical_pedagogy")
            cp_recommendations = extract_recommendations_from_response(cp_response)
            cp_strengths = extract_strengths_from_response(cp_response)
            cp_areas = extract_areas_for_improvement_from_response(cp_response)
            logger.info(f"Critical Pedagogy Score: {critical_pedagogy_score}/100")

            # ── AGENT 4: Lesson Design Quality ──
            logger.info(f"Agent 4/{provider.upper()}: Evaluating Lesson Design Quality...")
            ldq_prompt = gpt_design_prompt_template.format(lesson_plan_text=text)
            ldq_response = await llm_client.call(
                llm_name, ldq_prompt, agent_name=f"{provider.upper()}-Design"
            )

            lesson_design_score = extract_score_from_response(ldq_response, "lesson_design")
            ldq_recommendations = extract_recommendations_from_response(ldq_response)
            ldq_strengths = extract_strengths_from_response(ldq_response)
            ldq_areas = extract_areas_for_improvement_from_response(ldq_response)
            logger.info(f"Lesson Design Quality Score: {lesson_design_score}/100")

            # ── Build agent_responses IMMEDIATELY after all agents finish ──
            # BUG FIX: This was previously done AFTER the improvement generation
            # block, which meant the improvement block had an empty list to read from.
            agent_responses = [
                {
                    "agent": f"{provider.upper()}-PlaceBased",
                    "model": model_name,
                    "role": "Place-Based Learning Specialist",
                    "dimension": "place_based_learning",
                    "response": pbl_response[:1000],
                    "recommendations": pbl_recommendations[:5],
                    "score": place_based_score,
                    "analysis": {
                        "place_based_learning": {
                            "score": place_based_score,
                            "color": score_color(place_based_score),
                            "strengths": pbl_strengths[:5],
                            "areas_for_improvement": pbl_areas[:5],
                            "recommendations": pbl_recommendations[:5],
                        }
                    },
                },
                {
                    "agent": f"{provider.upper()}-Cultural",
                    "model": model_name,
                    "role": "Cultural Responsiveness Specialist",
                    # BUG FIX: use "cultural_responsiveness_integrated" consistently
                    "dimension": "cultural_responsiveness_integrated",
                    "response": crmp_response[:1000],
                    "recommendations": crmp_recommendations[:5],
                    "score": cultural_score,
                    "analysis": {
                        # BUG FIX: key must match frontend expectation
                        "cultural_responsiveness_integrated": {
                            "score": cultural_score,
                            "color": score_color(cultural_score),
                            "strengths": crmp_strengths[:5],
                            "areas_for_improvement": crmp_areas[:5],
                            "recommendations": crmp_recommendations[:5],
                        }
                    },
                },
                {
                    "agent": f"{provider.upper()}-Critical",
                    "model": model_name,
                    "role": "Critical Pedagogy Specialist",
                    "dimension": "critical_pedagogy",
                    "response": cp_response[:1000],
                    "recommendations": cp_recommendations[:5],
                    "score": critical_pedagogy_score,
                    "analysis": {
                        "critical_pedagogy": {
                            "score": critical_pedagogy_score,
                            "color": score_color(critical_pedagogy_score),
                            "strengths": cp_strengths[:5],
                            "areas_for_improvement": cp_areas[:5],
                            "recommendations": cp_recommendations[:5],
                        }
                    },
                },
                {
                    "agent": f"{provider.upper()}-Design",
                    "model": model_name,
                    "role": "Lesson Design Quality Specialist",
                    "dimension": "lesson_design_quality",
                    "response": ldq_response[:1000],
                    "recommendations": ldq_recommendations[:5],
                    "score": lesson_design_score,
                    "analysis": {
                        "lesson_design_quality": {
                            "score": lesson_design_score,
                            "color": score_color(lesson_design_score),
                            "strengths": ldq_strengths[:5],
                            "areas_for_improvement": ldq_areas[:5],
                            "recommendations": ldq_recommendations[:5],
                        }
                    },
                },
            ]

            # ── Merge recommendations ──
            recommendations = merge_and_deduplicate_recommendations(
                [
                    pbl_recommendations,
                    crmp_recommendations,
                    cp_recommendations,
                    ldq_recommendations,
                ],
                max_total=12,
            )
            logger.info(f"Total unique recommendations: {len(recommendations)}")

            # ── Compute composite score (dynamic weights) ──
            logger.info("Computing composite score (Framework v3.0)...")

            active_dimensions = {}
            if place_based_score > 0:
                active_dimensions["place_based_learning"] = place_based_score
            if cultural_score > 0:
                active_dimensions["cultural_responsiveness_integrated"] = cultural_score
            if critical_pedagogy_score > 0:
                active_dimensions["critical_pedagogy"] = critical_pedagogy_score
            if lesson_design_score > 0:
                active_dimensions["lesson_design_quality"] = lesson_design_score

            if active_dimensions:
                original_weights = framework_loader.get_scoring_weights()
                active_weights = {
                    k: original_weights.get(k, 0) for k in active_dimensions
                }
                total_weight = sum(active_weights.values())

                if total_weight > 0:
                    normalized_weights = {
                        k: v / total_weight for k, v in active_weights.items()
                    }
                    overall_score = calculate_weighted_score(
                        active_dimensions, normalized_weights
                    )
                else:
                    overall_score = sum(active_dimensions.values()) // len(
                        active_dimensions
                    )

                logger.info(f"Active dimensions: {list(active_dimensions.keys())}")
                for dim, sc in active_dimensions.items():
                    w = normalized_weights.get(dim, 0) if total_weight > 0 else 0
                    logger.info(f"  {dim}: {sc} (weight: {w * 100:.0f}%)")
                logger.info(f"Composite Score: {overall_score}/100")
            else:
                logger.warning("No valid scores from any API")
                overall_score = 0

            # ── Generate improved lesson plan if score is low ──
            if 0 < overall_score < 70:
                logger.info(
                    f"Score {overall_score} below 70, generating improved lesson plan..."
                )

                # BUG FIX: Now agent_responses is populated, so this extraction works
                all_strengths = []
                all_areas = []
                for agent_resp in agent_responses:
                    if agent_resp.get("analysis"):
                        for _dim_key, dim_data in agent_resp["analysis"].items():
                            if dim_data.get("strengths"):
                                all_strengths.extend(dim_data["strengths"][:3])
                            if dim_data.get("areas_for_improvement"):
                                all_areas.extend(dim_data["areas_for_improvement"][:3])

                all_strengths = list(dict.fromkeys(all_strengths))[:8]
                all_areas = list(dict.fromkeys(all_areas))[:8]

                scores_dict = {
                    "place_based_learning": place_based_score,
                    "cultural_responsiveness_integrated": cultural_score,
                    "critical_pedagogy": critical_pedagogy_score,
                    "lesson_design_quality": lesson_design_score,
                    "overall": overall_score,
                }

                try:
                    improvement_prompt = _load_improvement_prompt(
                        title=title,
                        grade_level=grade_level or "Not specified",
                        subject_area=subject_area or "Not specified",
                        text=text,
                        scores=scores_dict,
                        all_strengths=all_strengths,
                        all_areas=all_areas,
                        recommendations=recommendations,
                    )

                    logger.info(
                        f"Sending improvement request to Claude ({len(improvement_prompt)} chars)..."
                    )

                    ai_response = await asyncio.wait_for(
                        llm_client.call("claude", improvement_prompt),
                        timeout=300,
                    )

                    logger.info(f"Received improvement response ({len(ai_response)} chars)")

                    # Validate response format
                    has_bad_format = (
                        "1.1" in ai_response
                        or "['Understanding" in ai_response
                        or '["' in ai_response[:500]
                    )

                    if has_bad_format:
                        logger.warning(
                            "Claude returned structured format, attempting retry..."
                        )
                        retry_prompt = (
                            "Your previous response was in WRONG FORMAT (Python lists/numbered sections).\n"
                            "Rewrite the entire lesson plan in flowing narrative paragraphs.\n\n"
                            + improvement_prompt
                        )
                        try:
                            ai_response = await asyncio.wait_for(
                                llm_client.call("claude", retry_prompt),
                                timeout=300,
                            )
                            logger.info(f"Retry response ({len(ai_response)} chars)")
                        except Exception as retry_err:
                            logger.error(f"Retry failed: {retry_err}, using original")

                    if ai_response and len(ai_response) > 500:
                        lesson_plan_text = ai_response.strip()
                        logger.info(
                            f"Generated improved lesson plan ({len(lesson_plan_text)} chars)"
                        )

                        validation = validate_lesson_format(lesson_plan_text)
                        if validation["valid"]:
                            logger.info("Lesson plan format validation PASSED")
                        else:
                            logger.warning(
                                f"Lesson plan format issues: {validation['issues']}"
                            )
                    else:
                        logger.warning(
                            f"Generated plan too short ({len(ai_response)} chars), using original"
                        )
                        lesson_plan_text = text

                except asyncio.TimeoutError:
                    logger.error("Lesson plan generation timeout after 300s")
                    lesson_plan_text = text
                except Exception as gen_err:
                    logger.error(f"Lesson plan generation error: {gen_err}")
                    traceback.print_exc()
                    lesson_plan_text = text
            else:
                if overall_score == 0:
                    logger.warning("No valid scores, skipping improvement generation")
                else:
                    logger.info(
                        f"Score {overall_score} >= 70, no auto-improvement needed"
                    )
                lesson_plan_text = text

        except asyncio.TimeoutError:
            logger.error("Overall evaluation timeout")
            raise HTTPException(status_code=504, detail="Evaluation timeout")
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Evaluation error: {e}")
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Evaluation failed: {e}")

    else:
        # ── Mock mode — Framework v3.0 ──
        logger.info("MOCK mode with Framework v3.0 data (4 dimensions)")

        place_based_score = 72
        cultural_score = 68
        critical_pedagogy_score = 75
        lesson_design_score = 78

        weights = framework_loader.get_scoring_weights()
        dimension_scores = {
            "place_based_learning": place_based_score,
            "cultural_responsiveness_integrated": cultural_score,
            "critical_pedagogy": critical_pedagogy_score,
            "lesson_design_quality": lesson_design_score,
        }
        overall_score = calculate_weighted_score(dimension_scores, weights)

        logger.info(
            f"Mock Scores: PBL={place_based_score}, CRMP={cultural_score}, "
            f"CP={critical_pedagogy_score}, LDQ={lesson_design_score}, Overall={overall_score}"
        )

        agent_responses = [
            {
                "agent": "DeepSeek",
                "model": "deepseek-chat",
                "role": "Place-based Learning Specialist",
                "dimension": "place_based_learning",
                "response": "Mock evaluation - Place-based learning analysis",
                "analysis": build_analysis_structure(
                    "place_based_learning",
                    place_based_score,
                    "Mock evaluation - Good foundation in place-based learning.",
                    recommendations=[
                        "Name specific local places (e.g., 'Waitemata Harbour')",
                        "Partner with named local organisations",
                        "Design concrete fieldwork activities with protocols",
                        "Connect to specific local environmental issues",
                    ],
                    strengths=[
                        "Uses examples from the local community context",
                        "Encourages outdoor learning activities",
                        "Connects content to regional geography and ecosystems",
                        "Shows awareness of place-based pedagogy principles",
                    ],
                    areas_for_improvement=[
                        "Local references are too generic",
                        "Community partnerships mentioned but not detailed",
                        "Fieldwork activities described vaguely",
                        "Missing integration of local iwi environmental knowledge",
                    ],
                    gaps=[
                        "No specific local iwi partnerships identified",
                        "Indigenous ecological knowledge not incorporated",
                        "Community engagement is not substantive",
                    ],
                ),
                "recommendations": [
                    "Name specific local places",
                    "Partner with named local organisations",
                    "Design concrete fieldwork activities",
                ],
                "score": place_based_score,
                "time": 0.5,
            },
            {
                "agent": "Claude",
                "model": "claude-sonnet-4-20250514",
                "role": "Cultural Responsiveness & Maori Perspectives Specialist (Integrated)",
                # BUG FIX: consistent key
                "dimension": "cultural_responsiveness_integrated",
                "response": "Mock evaluation - Integrated cultural responsiveness analysis",
                "analysis": build_analysis_structure(
                    "cultural_responsiveness_integrated",
                    cultural_score,
                    "Mock evaluation - Demonstrates cultural awareness.",
                    recommendations=[
                        "Include more Te Reo Maori vocabulary",
                        "Consult with local iwi for cultural protocols",
                        "Add culturally diverse perspectives",
                        "Embed matauranga Maori more deeply",
                        "Use appropriate karakia and mihi protocols",
                    ],
                    strengths=[
                        "Acknowledges cultural context",
                        "References tikanga Maori appropriately",
                        "Creates inclusive learning environment",
                        "Shows respect for Maori worldviews",
                        "Includes some basic Te Reo Maori vocabulary",
                    ],
                    areas_for_improvement=[
                        "Te Reo usage limited to basic greetings",
                        "Matauranga Maori is peripheral rather than central",
                        "Lacks depth in Maori philosophical concepts",
                        "No evidence of iwi/hapu consultation",
                        "Cultural diversity acknowledged but not deeply woven in",
                        "Tikanga protocols not explicitly included",
                    ],
                    gaps=[
                        "No consultation with local iwi mentioned",
                        "Missing karakia or cultural protocols",
                        "Limited depth in Maori philosophical concepts",
                    ],
                    cultural_elements=[
                        "Basic Te Reo vocabulary: kia ora, ka pai, whanau",
                        "General tikanga references (not specific)",
                        "Acknowledgment of local iwi (name not specified)",
                        "Mention of cultural diversity (limited detail)",
                    ],
                ),
                "recommendations": [
                    "Include more Te Reo Maori vocabulary",
                    "Consult with local iwi for cultural validation",
                    "Embed specific matauranga Maori concepts",
                    "Include a whakatauki relevant to lesson topic",
                    "Add explicit karakia or mihi protocols",
                ],
                "score": cultural_score,
                "time": 0.6,
            },
            {
                "agent": "GPT-Critical",
                "model": "gpt-4o",
                "role": "Critical Pedagogy Specialist",
                "dimension": "critical_pedagogy",
                "response": "Mock evaluation - Critical pedagogy analysis",
                "analysis": build_analysis_structure(
                    "critical_pedagogy",
                    critical_pedagogy_score,
                    "Mock evaluation - Good pedagogical strategies.",
                    recommendations=[
                        "Add explicit critical questions challenging dominant narratives",
                        "Provide genuine student choice in topics and formats",
                        "Include formative assessment with self-reflection prompts",
                        "Design collaborative problem-solving tasks",
                        "Incorporate diverse voices and perspectives",
                    ],
                    strengths=[
                        "Uses questioning strategies for critical thinking",
                        "Includes student discussion opportunities",
                        "Offers some choice in learning activities",
                        "Shows awareness of active learning principles",
                        "Attempts meaningful dialogue",
                    ],
                    areas_for_improvement=[
                        "Student voice in decision-making is limited",
                        "Critical analysis of power structures is minimal",
                        "Assessment is mostly teacher-directed",
                        "Social justice themes are absent or superficial",
                        "Questioning is mostly closed-ended",
                        "Student choice is token rather than substantive",
                    ],
                    gaps=[
                        "No explicit examination of power dynamics",
                        "Missing critical questions about whose knowledge is valued",
                        "Limited connection to students' lived experiences",
                        "No action-oriented component (praxis)",
                    ],
                ),
                "recommendations": [
                    "Add explicit critical questions",
                    "Provide genuine student choice",
                    "Include formative assessment with self-reflection",
                    "Design collaborative problem-solving tasks",
                ],
                "score": critical_pedagogy_score,
                "time": 0.4,
            },
            {
                "agent": "GPT-Design",
                "model": "gpt-4o",
                "role": "Lesson Design & Quality Specialist",
                "dimension": "lesson_design_quality",
                "response": "Mock evaluation - Lesson design quality analysis (Framework v3.0)",
                "analysis": build_analysis_structure(
                    "lesson_design_quality",
                    lesson_design_score,
                    "Mock evaluation - Good instructional design quality.",
                    recommendations=[
                        "Develop detailed rubrics with clear success criteria",
                        "Add explicit differentiation strategies",
                        "Strengthen alignment between objectives and assessment",
                        "Improve time allocation realism",
                        "Provide more specific instructional guidance",
                    ],
                    strengths=[
                        "Learning objectives are clearly stated and measurable",
                        "Logical flow and coherent structure",
                        "Uses varied instructional strategies",
                        "Identifies appropriate resources and materials",
                        "Shows understanding of backward design principles",
                    ],
                    areas_for_improvement=[
                        "Assessment criteria are vague - lacks rubrics",
                        "Differentiation is limited",
                        "Time estimates may be unrealistic",
                        "Alignment between objectives and assessments could be stronger",
                        "Instructional guidance lacks detail",
                        "Limited variety in assessment methods",
                    ],
                    gaps=[
                        "No explicit rubrics or scoring guides",
                        "Missing scaffolds for struggling learners",
                        "Extensions for advanced learners not specified",
                        "Transitions between activities not detailed",
                    ],
                ),
                "recommendations": [
                    "Develop detailed rubrics",
                    "Add differentiation strategies",
                    "Strengthen objective-assessment alignment",
                    "Improve time allocation",
                ],
                "score": lesson_design_score,
                "time": 0.3,
            },
        ]

        recommendations = [
            "Strengthen local context integration with specific regional examples",
            "Include more Te Reo Maori vocabulary throughout",
            "Enhance student agency with more choice in learning activities",
            "Add community partnerships with local organisations",
            "Consult with local iwi for cultural protocols",
            "Embed matauranga Maori knowledge systems more deeply",
            "Add explicit rubrics with clear success criteria",
            "Include fieldwork opportunities in nearby environments",
            "Use appropriate karakia and mihi protocols",
            "Strengthen alignment between objectives and assessments",
            "Add formative assessment with student self-reflection",
            "Include more differentiation strategies for diverse learners",
        ]

        lesson_plan_text = text

    # ── Save to database ──
    try:
        logger.info("Saving evaluation to database...")

        eval_id = db.create_evaluation(
            lesson_plan_text=text,
            lesson_plan_title=title,
            grade_level=grade_level,
            subject_area=subject_area,
            api_mode=API_MODE,
            provider=provider,
        )

        db.update_evaluation_scores(
            eval_id=eval_id,
            place_based_score=place_based_score,
            cultural_score=cultural_score,
            overall_score=overall_score,
        )

        db.update_evaluation_results(
            eval_id=eval_id,
            agent_responses=agent_responses,
            debate_transcript={},
            recommendations=recommendations,
            status="completed",
        )

        logger.info(f"Evaluation saved (ID: {eval_id})")

    except Exception as db_err:
        logger.error(f"Database error: {db_err}")
        traceback.print_exc()

    # ── Return results ──
    framework_meta = framework_loader.load_theoretical_framework().get(
        "framework_metadata", {}
    )

    return {
        "status": "success",
        "evaluation_id": eval_id,
        "agent_responses": agent_responses,
        "recommendations": recommendations,
        "scores": {
            "place_based_learning": place_based_score,
            "cultural_responsiveness_integrated": cultural_score,
            "critical_pedagogy": critical_pedagogy_score,
            "lesson_design_quality": lesson_design_score,
            "overall": overall_score,
        },
        "framework_info": {
            "weights_applied": framework_loader.get_scoring_weights(),
            "dimensions_evaluated": [
                "place_based_learning",
                "cultural_responsiveness_integrated",
                "critical_pedagogy",
                "lesson_design_quality",
            ],
            "framework_version": framework_meta.get("version", "3.0"),
            "apis_used": {
                "deepseek": ENABLE_DEEPSEEK and place_based_score > 0,
                "claude": ENABLE_CLAUDE and cultural_score > 0,
                "gpt": ENABLE_GPT
                and (critical_pedagogy_score > 0 or lesson_design_score > 0),
            },
        },
        "improved_lesson_plan": lesson_plan_text,
        "mode": API_MODE,
    }

# ============================================================
# Debate Endpoint
# ============================================================
@app.post("/api/evaluate/lesson-debate")
async def evaluate_lesson_with_debate(
    request: EvaluationCreate,
    db: Database = Depends(get_db),
):
    """
    Evaluate with full multi-agent debate (3 rounds).
    Phase 1: Standard independent evaluation (reuses existing logic)
    Phase 2: Cross-review debate
    Phase 3: Consensus building
    """
    try:
        logger.info("=" * 60)
        logger.info("DEBATE MODE: Starting multi-agent debate evaluation")
        logger.info("=" * 60)

        # ── Phase 1: Run standard evaluation first ──
        logger.info("Phase 1: Running independent evaluations...")
        standard_result = await evaluate_lesson_plan(request, db)

        agent_responses = standard_result.get("agent_responses", [])

        if len(agent_responses) < 2:
            logger.warning("Only 1 agent responded, skipping debate")
            standard_result["evaluation_mode"] = "standard (insufficient agents for debate)"
            return standard_result

        # ── Phase 2 & 3: Run debate ──
        logger.info(f"Phase 2-3: Starting debate with {len(agent_responses)} agents...")

        from app.services.debate_engine import DebateEngine
        debate_engine = DebateEngine()

        debate_result = await debate_engine.run_debate(
            initial_evaluations=agent_responses,
            lesson_plan=request.lesson_plan_text,
            lesson_title=request.lesson_plan_title or "Untitled",
        )

        # ── Merge consensus scores into result ──
        consensus = debate_result.get("consensus", {})
        if consensus and "consensus_scores" in consensus:
            old_scores = standard_result.get("scores", {})
            new_scores = consensus["consensus_scores"]
            logger.info(f"Score changes after debate:")
            for dim in ["place_based_learning", "cultural_responsiveness_integrated",
                        "critical_pedagogy", "lesson_design_quality", "overall"]:
                old = old_scores.get(dim, "N/A")
                new = new_scores.get(dim, "N/A")
                if old != new:
                    logger.info(f"  {dim}: {old} → {new}")
            standard_result["scores"] = new_scores

        standard_result["debate_transcript"] = debate_result
        standard_result["evaluation_mode"] = "multi-agent-debate"

        logger.info("DEBATE MODE: Complete")
        logger.info("=" * 60)

        return standard_result

    except Exception as e:
        logger.error(f"Debate evaluation error: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Debate evaluation failed: {e}",
        )

# ============================================================
# Improve Lesson Endpoint
# ============================================================
@app.post("/api/improve-lesson")
@timing_decorator
async def improve_lesson(request: ImproveLessonRequest):
    """
    Generate an improved lesson plan via Claude — Framework v3.0.
    Produces a narrative-style lesson plan (no template filling).
    """
    try:
        logger.info(f"Improve Lesson Request: {request.lesson_title}")
        logger.info(f"Grade: {request.grade_level}, Recommendations: {len(request.recommendations)}")

        recs_text = "\n".join(
            f"{i + 1}. {rec}" for i, rec in enumerate(request.recommendations[:10])
        )

        improvement_prompt = f"""You are a highly experienced educator in Aotearoa New Zealand.
Your task is to significantly improve this lesson plan based on evaluation feedback.

CRITICAL OUTPUT FORMAT REQUIREMENT:
- Write in flowing narrative paragraphs like a real lesson plan.
- DO NOT use Python lists, numbered sections (1.1, 1.2), JSON, or dictionary format.
- Write naturally as if you're a teacher writing for other teachers.
- Start with: **IMPROVED LESSON PLAN: {request.lesson_title.upper()}**

CONTEXT:
Title: {request.lesson_title}
Grade Level: {request.grade_level or 'Not specified'}
Subject: {request.subject_area or 'Not specified'}

Current Scores:
{json.dumps(request.scores, indent=2)}

Key Recommendations to Address:
{recs_text}

Original Lesson Plan (excerpt):
{request.original_lesson[:3000]}

REQUIREMENTS:
1. Integrate 3+ Te Reo Maori terms with correct macrons (a, e, i, o, u)
2. Name 1-2 specific local places (e.g., Orakei Marae, Auckland Museum, Te Papa)
3. Include 1-2 critical thinking questions ("Whose stories do we usually hear?")
4. Provide 2-3 activities with clear steps in narrative form
5. Include an assessment rubric (table format OK for rubric only)
6. List 5-8 specific resources with titles and URLs
7. Target 1200-2000 words total

Include sections: Overview, Learning Objectives, Cultural Preparation,
Lesson Activities (2-3 activities), Assessment with Rubric, Resources, Time Allocation,
and a Note to Teachers about cultural consultation.

End with a cultural disclaimer:
"This AI-generated lesson plan requires review and adaptation before use.
Please consult with local iwi and cultural advisors for Maori content."

NOW BEGIN WRITING THE IMPROVED LESSON PLAN:
"""

        logger.info(f"Sending to Claude ({len(improvement_prompt)} chars)...")

        llm_client = LLMClient()
        response = await asyncio.wait_for(
            llm_client.call("claude", improvement_prompt),
            timeout=300,
        )

        logger.info(f"Received Claude response ({len(response)} chars)")

        improved_lesson = response.strip()

        # Validate format
        validation = validate_lesson_format(improved_lesson)
        if validation["valid"]:
            logger.info(
                f"Validation PASSED: {validation['word_count']} words, "
                f"Te Reo={validation['has_te_reo']}, "
                f"Places={validation['has_specific_places']}"
            )
        else:
            logger.warning(f"Validation issues: {validation['issues']}")

        if validation["warnings"]:
            logger.warning(f"Validation warnings: {validation['warnings']}")

        return {
            "status": "success",
            "improved_lesson": improved_lesson,
            "original_title": request.lesson_title,
            "recommendations_applied": len(request.recommendations),
            "framework_version": "3.0",
            "generator": "claude",
            "word_count": len(improved_lesson.split()),
        }

    except asyncio.TimeoutError:
        raise HTTPException(
            status_code=504,
            detail="Claude improvement generation timeout after 300s",
        )
    except Exception as e:
        logger.error(f"Error improving lesson: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Improvement failed: {e}")


# ============================================================
# Convert to Word Document
# ============================================================
@app.post("/api/convert-to-word")
async def convert_to_word(request: ConvertToWordRequest):
    """Convert lesson plan text to a downloadable Word document."""
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading(request.title or "Improved Lesson Plan", 0)

        for paragraph in request.content.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph)

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        filename = request.filename or "Improved_Lesson_Plan.docx"

        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    except ImportError:
        raise HTTPException(
            status_code=501,
            detail="python-docx is required for Word conversion but is not installed.",
        )
    except Exception as e:
        logger.error(f"Word conversion error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================
# Evaluations CRUD
# ============================================================
@app.get("/api/evaluations")
async def get_evaluations(
    limit: int = 20,
    offset: int = 0,
    db: Database = Depends(get_db),
):
    """List saved evaluation records."""
    try:
        logger.info(f"Fetching evaluations: limit={limit}, offset={offset}")

        evaluations = db.get_all_evaluations(limit=limit)

        if not evaluations:
            logger.info("No evaluations found")
            return {"status": "success", "evaluations": [], "count": 0}

        logger.info(f"Retrieved {len(evaluations)} evaluations")

        formatted = []
        for rec in evaluations:
            try:
                _agent_resp = json.loads(rec.get("agent_responses", "[]"))
            except (json.JSONDecodeError, TypeError):
                _agent_resp = []

            scores = {
                "place_based_learning": rec.get("place_based_score", 0),
                "cultural_responsiveness_integrated": rec.get("cultural_score", 0),
                "critical_pedagogy": rec.get("critical_pedagogy_score", 0),
                "lesson_design_quality": rec.get("lesson_design_score", 0),
            }

            formatted.append(
                {
                    "id": rec.get("id"),
                    # BUG FIX: frontend History reads "lesson_plan_title"
                    "lesson_plan_title": rec.get("lesson_plan_title")
                    or rec.get("lesson_title", "Untitled"),
                    "grade_level": rec.get("grade_level", "N/A"),
                    "subject_area": rec.get("subject_area", "N/A"),
                    "overall_score": rec.get("overall_score", 0),
                    "scores": scores,
                    "created_at": rec.get("created_at"),
                    "status": rec.get("status", "completed"),
                    "mode": rec.get("api_mode", "real"),
                    "framework_version": "3.0",
                }
            )

        return {"status": "success", "evaluations": formatted, "count": len(formatted)}

    except Exception as e:
        logger.error(f"Error fetching evaluations: {e}")
        traceback.print_exc()
        return {"status": "success", "evaluations": [], "count": 0}


@app.get("/api/evaluations/{evaluation_id}")
async def get_evaluation_by_id(
    evaluation_id: int,
    db: Database = Depends(get_db),
):
    """Get detailed information for a single evaluation record."""
    try:
        logger.info(f"Fetching evaluation ID: {evaluation_id}")

        evaluation = db.get_evaluation(evaluation_id)

        if not evaluation:
            raise HTTPException(
                status_code=404,
                detail=f"Evaluation {evaluation_id} not found",
            )

        logger.info(
            f"Retrieved evaluation: {evaluation.get('lesson_plan_title', 'Untitled')}"
        )

        # Parse JSON fields safely
        try:
            evaluation["agent_responses"] = json.loads(
                evaluation.get("agent_responses", "[]")
            )
        except (json.JSONDecodeError, TypeError):
            evaluation["agent_responses"] = []

        try:
            evaluation["recommendations"] = json.loads(
                evaluation.get("recommendations", "[]")
            )
        except (json.JSONDecodeError, TypeError):
            evaluation["recommendations"] = []

        try:
            evaluation["debate_transcript"] = json.loads(
                evaluation.get("debate_transcript", "{}")
            )
        except (json.JSONDecodeError, TypeError):
            evaluation["debate_transcript"] = {}

        return {
            "status": "success",
            "evaluation": evaluation,
            "framework_version": "3.0",
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching evaluation: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Failed to retrieve evaluation: {e}",
        )


@app.delete("/api/evaluations/{evaluation_id}")
async def delete_evaluation(
    evaluation_id: int,
    db: Database = Depends(get_db),
):
    """Delete a specific evaluation record."""
    try:
        logger.info(f"Deleting evaluation ID: {evaluation_id}")

        evaluation = db.get_evaluation(evaluation_id)
        if not evaluation:
            raise HTTPException(
                status_code=404,
                detail=f"Evaluation {evaluation_id} not found",
            )

        db.delete_evaluation(evaluation_id)
        logger.info(f"Deleted evaluation ID: {evaluation_id}")

        return {
            "status": "success",
            "message": f"Evaluation {evaluation_id} deleted successfully",
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting evaluation: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Failed to delete evaluation: {e}",
        )


# ============================================================
# Startup Information
# ============================================================
logger.info("=" * 60)
logger.info("Lesson Plan Evaluator API - Framework v3.0")
logger.info("=" * 60)

try:
    framework = framework_loader.load_theoretical_framework()
    framework_name = framework.get("framework_metadata", {}).get("name", "Default")
    framework_version = framework.get("framework_metadata", {}).get("version", "3.0")
    dimensions = list(framework.get("dimensions", {}).keys())

    logger.info(f"Framework: {framework_name} (v{framework_version})")
    logger.info(f"Dimensions: {len(dimensions)}")
    for dim in dimensions:
        logger.info(f"  - {dim}")
    logger.info(f"API Mode: {API_MODE}")
    logger.info("Database: Connected")

    agent_design = framework_loader.load_agent_design()
    agents = agent_design.get("agents", {})
    logger.info(f"Configured Agents ({len(agents)}):")
    for _agent_id, agent_info in agents.items():
        name = agent_info["name"]
        is_enabled = (
            (name == "DeepSeek" and ENABLE_DEEPSEEK)
            or (name == "Claude" and ENABLE_CLAUDE)
            or (name in ("GPT-Critical", "GPT-Design") and ENABLE_GPT)
        )
        status = "ENABLED" if is_enabled else "DISABLED"
        logger.info(f"  [{status}] {name}: {agent_info['role']}")

    weights = framework_loader.get_scoring_weights()
    logger.info("Scoring Weights (v3.0):")
    for dim, weight in weights.items():
        logger.info(f"  - {dim}: {weight * 100:.0f}%")

    logger.info("API Configuration:")
    logger.info(f"  DeepSeek: {'Enabled' if ENABLE_DEEPSEEK else 'Disabled'}")
    logger.info(f"  Claude:   {'Enabled' if ENABLE_CLAUDE else 'Disabled'}")
    logger.info(f"  GPT:      {'Enabled' if ENABLE_GPT else 'Disabled'}")
    logger.info(f"  Timeout:  {API_TIMEOUT}s")
    logger.info(f"  Retries:  {API_MAX_RETRIES}")
    logger.info(f"  Delay:    {API_RETRY_DELAY}s")

except Exception as e:
    logger.warning(f"Could not load framework details: {e}")

logger.info("=" * 60)
logger.info("API is ready to accept requests!")
logger.info("Endpoints:")
logger.info("  GET  /                          - API info")
logger.info("  GET  /api/framework             - Framework details")
logger.info("  GET  /api/framework/dimension/X - Dimension details")
logger.info("  GET  /api/health                - Health check")
logger.info("  POST /api/extract-text          - Extract text from file")
logger.info("  POST /api/evaluate/lesson       - Evaluate lesson plan")
logger.info("  POST /api/improve-lesson        - Generate improved lesson")
logger.info("  POST /api/convert-to-word       - Convert to Word document")
logger.info("  GET  /api/evaluations           - Evaluation history")
logger.info("  GET  /api/evaluations/{id}      - Single evaluation")
logger.info("  DELETE /api/evaluations/{id}    - Delete evaluation")
logger.info("=" * 60)